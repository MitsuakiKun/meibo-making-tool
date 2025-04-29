import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_ALIGN_VERTICAL  
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from io import BytesIO

st.title("名簿 Word ファイル作成ツール")

uploaded_file = st.file_uploader("Excelファイルをアップロードしてください", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    df = df.sort_values(by=['班順'])

    # Word文書の作成（ここにあなたの処理を貼り付け、ファイル保存をBytesIOに変更）
    doc = Document()
    styles = doc.styles
    style = styles.add_style('Original-Style', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Arial"
    font.size = Pt(22)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'BIZ UDPゴシック')
    font.bold = True

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for _, group in df.groupby('班順'):
        group = group.sort_values(by=['班順'])
        group_name = group['班名'].values[0] 
        doc.add_heading(group_name, level=1).style = 'Original-Style'
        table = doc.add_table(rows=len(group)*2 , cols=6)
        table.style = 'TableGrid'

        # データ行の作成
        for i, (_, row) in enumerate(group.iterrows(), start=0):
            # 各行の高さを設定（例：1.5cm）
            table.rows[i * 2].height = Cm(1.2)
            table.rows[i * 2 + 1].height = Cm(1.2)

            # 高さの種類を「最小」ではなく「固定」にする場合（必要なら）
            table.rows[i * 2].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            table.rows[i * 2 + 1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            # 顔写真のセルを結合
            photo_cell = table.cell(i * 2, 0)
            photo_cell.merge(table.cell(i * 2 + 1, 0))
            
            # 名前とふりがなを一つのセルに表示
            name_furigana_cell = table.cell(i * 2, 1)
            name_furigana_text = f'{row["名前"]}（{row["ふりがな"]}）'
            if not pd.isna(name_furigana_text):  # セルがnullでないことを確認
                name_furigana_cell.text = name_furigana_text
            else:
                name_furigana_cell.text = " "  # セルがnullの場合、空白文字を挿入
            name_furigana_cell.merge(table.cell(i * 2, 4))

            # Apply the 'Original-Style' to the paragraph
            name_furigana_paragraph = name_furigana_cell.paragraphs[0]
            name_furigana_paragraph.style = 'Original-Style'
            # Set the font on the run within the paragraph
            name_furigana_run = name_furigana_paragraph.runs[0]
            name_furigana_run.font.name = "BIZ UDPゴシック"
            name_furigana_run.font.size = Pt(22)
            name_furigana_run.font.bold = True

            # Apply the 'Original-Style' to other cells in the row
            gakusha_cell = table.cell(i * 2 , 5)
            gakusha_cell_text = row['学舎']
            if not pd.isna(gakusha_cell_text):  # セルがnullでないことを確認
                gakusha_cell.text = gakusha_cell_text
            else:
                gakusha_cell.text = " "  # セルがnullの場合、空白文字を挿入
            
            gakusha_cell.paragraphs[0].style = 'Original-Style'
            gakusha_run = gakusha_cell.paragraphs[0].runs[0]
            gakusha_run.font.name = "BIZ UDPゴシック"
            gakusha_run.font.size = Pt(20)
            gakusha_run.font.bold = True 

            university_cell = table.cell(i * 2 + 1, 1)
            university_cell_text = row['大学']
            if not pd.isna(university_cell_text):  # セルがnullでないことを確認
                university_cell.text = university_cell_text
            else:
                university_cell.text = " "  # セルがnullの場合、空白文字を挿入
            university_cell.paragraphs[0].style = 'Original-Style'
            university_cell.merge(table.cell(i * 2 + 1, 3))
            university_run = university_cell.paragraphs[0].runs[0]
            university_run.font.name = "BIZ UDPゴシック"
            university_run.font.size = Pt(19)
            university_run.font.bold = True

            grade_cell = table.cell(i * 2 + 1, 4)
            grade_cell_text = str(row['学年'])
            if not pd.isna(grade_cell_text):  # セルがnullでないことを確認
                grade_cell.text =grade_cell_text
            else:
                grade_cell.text = " "  # セルがnullの場合、空白文字を挿入
            grade_cell.paragraphs[0].style = 'Original-Style'
            grade_run = grade_cell.paragraphs[0].runs[0]
            grade_run.font.name = "BIZ UDPゴシック"
            grade_run.font.size = Pt(22)
            grade_run.font.bold = True


            sp_cell = table.cell(i * 2 + 1, 5)
            sp_cell_text = row['SP']
            if not pd.isna(sp_cell_text):  # セルがnullでないことを確認
                sp_cell.text = sp_cell_text
            else:
                sp_cell.text = " "  # セルがnullの場合、空白文字を挿入
            sp_cell.paragraphs[0].style = 'Original-Style'
            sp_run = sp_cell.paragraphs[0].runs[0]
            sp_run.font.name = "BIZ UDPゴシック"
            sp_run.font.size = Pt(22)
            sp_run.font.bold = True

            # 性別に応じて名前の文字色を設定
            if row['性別'] == '男':
                name_furigana_run.font.color.rgb = RGBColor(0, 0, 255)  # 青色（男性）
            elif row['性別'] == '女':
                name_furigana_run.font.color.rgb = RGBColor(255, 0, 0)  # 赤色（女性）
            
            name_furigana_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            gakusha_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            university_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            grade_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            sp_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER    # セルの中央揃え

    # 保存用バッファ
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("Wordファイルの作成が完了しました！")
    st.download_button(label="Wordファイルをダウンロード", data=buffer, file_name="名簿.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
